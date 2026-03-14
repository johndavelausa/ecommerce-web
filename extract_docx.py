from docx import Document
from pathlib import Path

def extract(path: str) -> str:
    doc = Document(path)
    out = []
    for p in doc.paragraphs:
        t = (p.text or '').strip()
        if t:
            out.append(t)
    for ti, table in enumerate(doc.tables):
        out.append(f"\n[TABLE {ti+1}]")
        for row in table.rows:
            cells = [(c.text or '').strip().replace('\n',' ') for c in row.cells]
            out.append(' | '.join(cells))
    return '\n'.join(out) + '\n'

files = ['ThriftStore_System_Requirements.docx','ThriftStore_Build_Order.docx']
for f in files:
    txt = extract(f)
    out_path = Path(f).with_suffix('.extracted.txt')
    out_path.write_text(txt, encoding='utf-8')
    print(f"wrote {out_path} ({len(txt)} chars)")
